import fitz
import io
import shutil
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from PIL import Image
import pytesseract
from docx import Document


# =====================================================
# CONFIG
# =====================================================

MIN_TEXT_THRESHOLD = 40
OCR_DPI = 300
MAX_WORKERS = 4   # increase if CPU is strong (6–8 ideal)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s — %(levelname)s — %(message)s"
)

logger = logging.getLogger("loader")


# =====================================================
# TESSERACT AUTO-DETECT
# =====================================================

tesseract_path = shutil.which("tesseract")

if tesseract_path is None:
    # Windows fallback
    pytesseract.pytesseract.tesseract_cmd = (
        r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    )
else:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path


# =====================================================
# TEXT NORMALIZER
# =====================================================

def normalize(text: str) -> str:
    """
    Clean text for better comparison accuracy.
    """

    if not text:
        return ""

    text = text.replace("\n", " ")
    text = text.replace("\t", " ")

    # collapse spaces
    text = " ".join(text.split())

    return text.strip()


# =====================================================
# OCR ENGINE
# =====================================================

def ocr_page(page):
    """
    Convert PDF page → image → OCR
    """

    pix = page.get_pixmap(dpi=OCR_DPI)

    img = Image.open(io.BytesIO(pix.tobytes("png")))

    text = pytesseract.image_to_string(
        img,
        config="--oem 3 --psm 6"
    )

    return normalize(text)


# =====================================================
# PAGE PROCESSOR (PARALLEL READY)
# =====================================================

def process_page(doc, page_number):
    """
    Extract text from a single page.
    Designed for parallel execution.
    """

    page = doc.load_page(page_number)

    blocks = []
    data = page.get_text("dict")

    for block in data["blocks"]:
        if "lines" in block:
            for line in block["lines"]:
                blocks.append(
                    " ".join(span["text"] for span in line["spans"])
                )

    text = normalize(" ".join(blocks))

    # OCR fallback
    if len(text) < MIN_TEXT_THRESHOLD:
        logger.info(f"OCR triggered → page {page_number+1}")
        text = ocr_page(page)

    return page_number + 1, text


# =====================================================
# FAST PDF READER (PARALLEL)
# =====================================================

def read_pdf(pdf_path):

    logger.info("Opening PDF...")

    doc = fitz.open(pdf_path)
    total_pages = len(doc)

    logger.info(f"Total pages: {total_pages}")
    logger.info(f"Using {MAX_WORKERS} CPU threads")

    pages = {}

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:

        futures = [
            executor.submit(process_page, doc, i)
            for i in range(total_pages)
        ]

        for future in as_completed(futures):
            page_num, text = future.result()
            pages[page_num] = text

    doc.close()

    # ensure correct order
    return dict(sorted(pages.items()))


# =====================================================
# DOCX READER
# =====================================================

def read_docx(docx_path):

    logger.info("Reading DOCX...")

    doc = Document(docx_path)

    content = []

    # paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)

    # tables
    for table in doc.tables:
        for row in table.rows:
            content.append(
                " | ".join(cell.text.strip() for cell in row.cells)
            )

    return {1: normalize(" ".join(content))}


# =====================================================
# MASTER LOADER
# =====================================================

def load_document(path):

    logger.info(f"Loading document: {path}")

    lower = path.lower()

    if lower.endswith(".pdf"):
        return read_pdf(path)

    elif lower.endswith(".docx"):
        return read_docx(path)

    else:
        raise ValueError("Unsupported file type. Only PDF/DOCX allowed.")
